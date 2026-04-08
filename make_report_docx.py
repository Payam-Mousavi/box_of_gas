#!/usr/bin/env python3
"""Convert REPORT.md into a formatted Word document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = Path(__file__).resolve().parent
REPORT_MD = HERE / "REPORT.md"
PLOTS_DIR = HERE / "report_plots"
OUTPUT = HERE / "report.docx"


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_image(doc, alt_text, img_path):
    """Insert an image centered, with a caption below."""
    real_path = PLOTS_DIR / Path(img_path).name
    if not real_path.exists():
        p = doc.add_paragraph(f"[Image not found: {img_path}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(real_path), width=Inches(5.5))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(alt_text)
    set_run_font(r, size=9, italic=True, color=(100, 100, 100))


def add_table(doc, header_row, data_rows):
    """Insert a styled table."""
    cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=cols)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, cell_text in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    # Data
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text
    doc.add_paragraph()  # spacing after table


def parse_table(lines, start_idx):
    """Parse a markdown table starting at start_idx. Returns (header, rows, end_idx)."""
    header_line = lines[start_idx].strip()
    cols = [c.strip().strip("*") for c in header_line.split("|")[1:-1]]
    # skip separator line
    idx = start_idx + 2
    rows = []
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [c.strip().strip("*") for c in lines[idx].split("|")[1:-1]]
        rows.append(cells)
        idx += 1
    return cols, rows, idx


def add_formatted_paragraph(doc, text, bold_whole=False):
    """Add a paragraph with inline **bold** and *italic* formatting."""
    p = doc.add_paragraph()
    # Split on bold and italic markers
    # Process **bold** first, then *italic*
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            # Check for italic inside bold
            sub_parts = re.split(r"(\*.*?\*)", inner)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*"):
                    r = p.add_run(sp[1:-1])
                    set_run_font(r, bold=True, italic=True)
                else:
                    r = p.add_run(sp)
                    set_run_font(r, bold=True)
        else:
            # Handle italic in non-bold text
            sub_parts = re.split(r"(\*.*?\*)", part)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*"):
                    r = p.add_run(sp[1:-1])
                    set_run_font(r, italic=True)
                else:
                    p.add_run(sp)
    return p


def convert():
    text = REPORT_MD.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # Strip markdown bold from headings
            title = title.replace("**", "")
            if level == 1:
                add_heading(doc, title, 0)
            else:
                add_heading(doc, title, level - 1)
            i += 1
            continue

        # Images
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2)
            add_image(doc, alt, path)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"\|[\s\-:|]+\|", lines[i + 1].strip()):
            header, rows, end_idx = parse_table(lines, i)
            add_table(doc, header, rows)
            i = end_idx
            continue

        # Bullet points
        bullet_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if bullet_match:
            add_formatted_paragraph(doc, bullet_match.group(2))
            p = doc.paragraphs[-1]
            p.style = "List Number"
            i += 1
            continue

        dash_match = re.match(r"^[-*]\s+(.*)", stripped)
        if dash_match:
            add_formatted_paragraph(doc, dash_match.group(1))
            p = doc.paragraphs[-1]
            p.style = "List Bullet"
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            if not next_stripped:
                break
            if next_stripped.startswith("#") or next_stripped.startswith("|") or next_stripped.startswith("!"):
                break
            if re.match(r"^[-*]\s+", next_stripped) or re.match(r"^\d+\.\s+", next_stripped):
                break
            para_lines.append(next_stripped)
            i += 1

        full_text = " ".join(para_lines)
        add_formatted_paragraph(doc, full_text)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    convert()
